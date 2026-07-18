import io
from pypdf import PdfReader
from docx import Document
from PIL import Image
from google import genai
from app.core.config import settings

def parse_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file bytes."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        # Fallback to latin-1 if utf-8 fails
        return file_bytes.decode("latin-1")

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes using PyPDF."""
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    extracted_text = []
    
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            extracted_text.append(text)
            
    return "\n\n--- Page Break ---\n\n".join(extracted_text)

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from Microsoft Word (.docx) file bytes."""
    docx_file = io.BytesIO(file_bytes)
    doc = Document(docx_file)
    extracted_text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text:
            extracted_text.append(paragraph.text)
            
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                extracted_text.append(" | ".join(row_text))
                
    return "\n".join(extracted_text)

def parse_image_ocr(file_bytes: bytes) -> str:
    """
    Extract text from images (PNG, JPG, JPEG) using Google Gemini's
    multimodal vision capabilities. High accuracy and format preservation.
    """
    if not settings.GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY environment variable is not configured.")
        
    try:
        # Initialize Gemini genai client
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        
        # Load image via Pillow
        image = Image.open(io.BytesIO(file_bytes))
        
        # Call multimodal Gemini 2.5 Flash model
        prompt = (
            "You are an expert OCR engine. Extract all text from this image "
            "as accurately as possible. Retain the original layout, formatting, "
            "and column structure where applicable. "
            "Do not add any preamble, comments, or extra text. Output ONLY the extracted text."
        )
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt, image]
        )
        
        if response.text:
            return response.text.strip()
        else:
            return "No text could be extracted from this image."
            
    except Exception as e:
        raise RuntimeError(f"OCR parsing failed via Gemini API: {str(e)}")

def extract_document_text(file_bytes: bytes, file_type: str) -> str:
    """
    Dispatches file bytes to the appropriate parser based on the lowercase file extension.
    """
    ftype = file_type.lower().strip(".")
    
    if ftype in ["txt", "text"]:
        return parse_txt(file_bytes)
    elif ftype == "pdf":
        return parse_pdf(file_bytes)
    elif ftype in ["docx", "doc"]:
        return parse_docx(file_bytes)
    elif ftype in ["png", "jpg", "jpeg"]:
        return parse_image_ocr(file_bytes)
    else:
        raise ValueError(f"Unsupported file type for parsing: {file_type}")
